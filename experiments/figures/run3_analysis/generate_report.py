"""Generate the experiment report as .docx file."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = OUT_DIR
REPO_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(OUT_DIR)))

doc = Document()

# ── Style Setup ───────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    if level == 1:
        heading_style.font.size = Pt(18)
    elif level == 2:
        heading_style.font.size = Pt(14)
    else:
        heading_style.font.size = Pt(12)

def add_para(text, bold=False, italic=False, size=None, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if align is not None: p.alignment = align
    return p

def add_figure(path, caption, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_para(caption, italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=(0x66, 0x66, 0x66))
    else:
        add_para(f'[Figure not found: {path}]', italic=True, size=9)

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

# ═══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('AI Racer 自动驾驶控制器', bold=True, size=24,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('算法实验与参赛报告', bold=True, size=18,
         align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('PKU DSAI Racer Competition — Complex Track', size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x55, 0x55, 0x55))
add_para('2026 年 6 月', size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
         color=(0x55, 0x55, 0x55))
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# CHAPTER 1: 算法思想阐述
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('第1章 算法思想阐述', level=1)

doc.add_heading('1.1 问题定义与整体方案', level=2)

add_para(
    'AI Racer 是一个基于视觉的自动驾驶竞赛平台。'
    '每辆赛车仅配备左右两个前向摄像头（各 640×480 像素），'
    '控制器的输入是这两张图像和当前仿真时间戳，'
    '输出是两个浮点数：steering（方向盘转角，范围 [-1, 1]，负为左转）和 speed（车速比例，范围 [0, 1]）。'
    '赛车的物理模型包含真实的质量、轮胎摩擦和碰撞动力学，因此控制器必须能够处理传感器噪声、'
    '道路纹理变化、其他车辆遮挡以及碰撞后的恢复。'
)

add_para(
    '我们的方案采用经典的"感知-估计-控制"三层流水线（Perception-Estimation-Control Pipeline）。'
    '感知层从原始图像中提取道路区域的几何描述；估计层将这些描述转化为连续的赛道状态量；'
    '控制策略层根据赛道状态计算转向和速度指令。三层之间通过严格定义的数据结构传递信息，'
    '彼此解耦，使得每一层可以独立测试和优化。'
    '图 1.1 展示了这一流水线的整体结构。'
)

add_figure(os.path.join(FIG_DIR, 'fig1_pipeline.png'),
           '图 1.1 控制流水线（Control Pipeline）')

doc.add_heading('1.2 视觉感知：从图像像素到道路模型', level=2)

add_para(
    '感知模块（perception.py）的输入是左右两张 BGR 图像，输出是一个 PerceptionObs 数据结构，'
    '包含道路中心线的采样点坐标、左右边界点、白线（车道线）的偏移与朝向、以及置信度和障碍物标志。'
    '核心任务是从像素中分离"路面"和"非路面"。'
)

add_para('1.2.1 颜色空间多通道路面分割', bold=True, size=11)

add_para(
    '道路在图像中呈现为深灰色沥青，但光照变化、阴影和不同赛道材质（complex 赛道的红色地面、'
    'basic 赛道的灰色沥青）使单一颜色阈值不可靠。我们采用了多通道联合判据：'
    '对每个像素同时检查 HSV（色相、饱和度、亮度）和 LAB（CIE L*a*b*）两个颜色空间，以及通道差分条件'
    '（B-G 和 G-R 的差值范围）。只有同时满足所有通道条件的像素才被归类为道路。'
    '这种方式比单一阈值更鲁棒：例如，草地（绿色的高饱和度）在 HSV 的饱和度通道会被排除，'
    '红色地面在 LAB 的 a* 通道（红-绿轴）会有显著不同的特征，因此能被单独的颜色卡区分开来。'
)

add_para(
    '车道线（白色虚线）的检测采用独立的颜色阈值和形态学处理。'
    '与道路分割不同，车道线需要更高的亮度和更低的色彩饱和度。'
    '我们还引入了"路面上下文验证"：一个白色像素要被认定为车道线，不仅自身要满足白色条件，'
    '其左右紧邻区域内的像素也必须属于深色路面。'
    '这有效排除了白色护栏（护栏外侧是草地或红色地面，而非路面）。'
)

add_para('1.2.2 扫描线道路中心提取', bold=True, size=11)

add_para(
    '得到道路的二值掩码（road mask）后，我们在图像的下半部分（ROI top=0.42，包含车前方最近的路面）'
    '水平扫描 12 条等间距的扫描线。每条扫描线上找到的连续道路段的中心点被记录为道路中心采样点。'
    '当道路在远处分叉或被其他车辆遮挡时，单条扫描线可能出现多个候选段。'
    '此时，我们利用上一帧保存的道路中心位置作为"时间锚点"'
    '（Temporal Anchor），选择与上帧中心最接近的候选段，从而保持道路追踪的时间连续性。'
    '这在 complex 赛道 CP3 区域（多个道路段在画面中合并）尤其关键。'
)

add_para('1.2.3 边界余量与障碍物检测', bold=True, size=11)

add_para(
    '除了道路中心点，感知模块还提取左右道路边界点，用于计算"边界余量"（margin）——'
    '即车身到最近左右边界的归一化距离。这一信息供后续策略层判断车辆是否贴近护栏。'
    '此外，独立的 opponent.py 模块在下半部中间 ROI 进行车身色块检测：利用亮白色和近黑色两种候选色，'
    '通过连通域面积和尺寸过滤噪声，判断近处是否有其他赛车。'
)

doc.add_heading('1.3 赛道几何估计：从离散点到连续状态', level=2)

add_para(
    '估计模块（estimator.py）将感知模块输出的离散中心点转化为连续的赛道几何状态 TrackState。'
    'TrackState 包含五个核心量，它们构成了控制策略的全部几何输入。'
)

add_para('1.3.1 五个核心几何量', bold=True, size=11)

add_para(
    'lateral_error（横向偏差）：车辆中心相对于道路目标线的横向偏移，单位为归一化比值（-1 到 1，'
    '负表示车辆偏左）。由近处（15% 道路进度）的中心点位置估算。'
)
add_para(
    'heading_error（航向偏差）：车辆当前朝向与道路目标方向的夹角偏差，单位为归一化比值。'
    '由 45% 道路进度处的中心点方向估算。'
)
add_para(
    'curvature（曲率）：道路前方的弯曲程度，通过对中心点做二次多项式拟合提取。'
    '曲率为负表示左弯，为正表示右弯。拟合需要至少 5 个有效采样点且纵向跨度足够。'
)
add_para(
    'lookahead_error（前瞻偏差）：较远处（75% 道路进度）的预测横向偏差，'
    '用于在车辆到达弯道之前就开始预判转向需求。'
)
add_para(
    'confidence（置信度）：0 到 1 之间的实数，衡量当前估计的可信程度。'
    '受采样点数量、道路宽度、掩码饱和度等因素影响。'
    '当 confidence 低于 lost_confidence（0.10）时，车辆判定为"丢线"（lost）状态。'
)

add_para('1.3.2 时序平滑与丢线衰减', bold=True, size=11)

add_para(
    '为防止单帧噪声导致控制抖动，所有几何估计量都经过指数滑动平均（EMA）平滑。'
    '当车辆进入丢线状态（前方无有效道路观测，如通过蓝门或天空区域）时，'
    '几何量不会立即跳变为 0，而是按照预设的衰减系数逐帧衰减。'
    '例如，横向偏差以 0.85 的系数衰减、航向偏差以 0.78 衰减。'
    '这种"惯性保持"策略让车辆在短暂的无观测期间能够维持原有航向滑行，直到道路重新出现。'
)

doc.add_heading('1.4 控制策略：从几何状态到驾驶决策', level=2)

add_para(
    '策略模块（policy.py）是控制系统的核心，它将 TrackState 转化为具体的 steering 和 speed 指令。'
    '整个策略围绕四个核心设计思想展开。'
)

add_para('1.4.1 多模式状态机', bold=True, size=11)

add_para(
    '策略不是用单一公式计算控制量，而是维护一个内部驾驶模式状态机，'
    '包含五种模式：cruise（巡航）、hard_turn（急弯）、correcting（回中修正）、'
    'lost（丢线）和 escaping（脱困）。模式切换由当前几何状态决定。'
    '例如，当 curve_risk（弯道风险，由曲率和航向偏差合成）超过 0.30 的阈值时，'
    '车辆进入 hard_turn 模式，此时转向会被放大（乘以 0.78 的缩放因子），'
    '速度上限被限制在 hard_turn_speed（0.72）。不同模式使用的转向平滑系数也不同，'
    '急弯模式使用较小的平滑系数以允许更快的转向响应。'
)

add_para('1.4.2 入弯时机门控（Turn-In Gate）', bold=True, size=11)

add_para(
    '这是整个控制策略中最重要的创新之一。问题背景：在直道接近弯道时，'
    '远处的前瞻项（lookahead_error）已经"看到"弯道并产生较大的预瞄转向需求，'
    '导致车辆在直道上就提前切向内线——即"入弯太早"（early turn-in），'
    '结果是在弯道 apex 处贴到内侧护栏。'
)

add_para(
    '我们的解决方案是引入"入弯时机门控"：用一个 0-1 之间的门控因子 corner_arrival 乘在远处预瞄项上。'
    '门控因子的计算仅基于近处横向偏差（lateral_error）：在直道接近段，车辆紧贴车道中心线，'
    '横向偏差接近 0，门控因子也接近 0，把远处预瞄项几乎完全压住，车辆继续直行。'
    '只有当车辆物理上到达弯道入口，开始偏离中心线时，横向偏差增大，门控因子随之打开，'
    '远处预瞄项才开始发挥作用。这个设计的关键洞察是：'
    '"车辆是否到达弯道"应该用近处的物理位置来判断（"我已经偏了吗？"），'
    '而不是用远处的视觉预判（"我看到了弯"）。'
)

add_para(
    '门控因子不是简单的瞬时计算，而是引入了跨帧 latch 保持机制：'
    '在 hard_turn 模式中，门控因子取最近峰值的最大值（ratchet），'
    '防止在弯中因道路掩码短暂回正导致门控关闭、车辆"转一半收轮"。'
    '离开弯道后，门控因子按 0.92 的衰减系数缓慢释放，提供出弯迟滞效果。'
    '此外，门控的参考基准随车速自动调整：高速时参考值缩小，门控提前打开，'
    '弥补高速下车辆在门控开启前就已冲过的距离。'
)

add_para('1.4.3 多层脱困系统（Self-Rescue Chain）', bold=True, size=11)

add_para(
    '在 6 车同场竞技的场景中，碰撞、挤压和卡栏杆是不可避免的极端情况。'
    '我们设计了一套四层递进的脱困系统，从最精确的几何脱困到完全不依赖几何的兜底安全网。'
    '图 1.2 展示了这四层的触发条件和动作。'
)

add_figure(os.path.join(FIG_DIR, 'fig3_escape.png'),
           '图 1.2 多层脱困系统（Multi-Layer Escape System）')

add_para(
    'Level 1（顶栏脱困 Pinned Escape，40 帧 ~1.3s）：当车辆横向偏差 ≥ 0.38 且当前转向 ≥ 0.48，'
    '表明车辆正以大舵角顶住栏杆，此时朝远离栏杆的方向施加 0.92 的硬舵。'
    'Level 2（边界障碍脱困 Boundary Obstacle，90 帧 ~2.9s）：当近处检测到障碍物且单侧边界余量 ≤ 0.08，'
    '朝开阔侧施加 0.92 的硬舵。Level 3（低速卡死死脱困 Low-Speed Stall，90 帧）：'
    '当指令速度降到 0.28 以下并持续 45 帧，触发带摆动的脱困——除了基准 0.95 的硬舵外，'
    '每 10 帧交替叠加 ±0.30 的摆动幅度。这种"摆动"模拟了人类司机在车轮卡死时的左右摇晃操作，'
    '能有效打破静摩擦。Level 4（强制脱困 Force Escape，120 帧 ~3.8s）：'
    '完全独立于几何条件，只要连续丢线 ≥ 45 帧或指令速度 ≤ 0.08 持续 60 帧就触发，'
    '是最后的安全底线。'
)

add_para('1.4.4 多车交互策略', bold=True, size=11)

add_para(
    '多车环境下的安全策略包含两个层面。第一层是速度控制：当 near_obstacle 标志为真时，'
    '目标速度首先乘以 0.72 的通用降速因子。如果在弯道中（curve_risk ≥ 0.25）同时检测到对手车，'
    '再额外乘以 0.55 的弯道叠加减速因子。这确保在 CP3 等急弯密集区域有多辆车聚集时，'
    '车辆以显著降低的速度通过，避免高速碰撞导致卡死。'
)

add_para(
    '第二层是主动避让转向：利用左右边界余量的差值来判断障碍物的方位——'
    '如果右侧余量明显小于左侧，说明障碍物在右侧，车辆向左施加一个最多 0.18 的舵角偏置。'
    '这个偏置叠加在正常转向输出上，形成"朝开阔空间偏转"的效果。'
)

# ═══════════════════════════════════════════════════════════════════
# CHAPTER 2: 程序代码说明
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('第2章 程序代码说明', level=1)

doc.add_heading('2.1 项目结构与模块职责', level=2)

add_para(
    '项目遵循模块化架构，controller/ 目录下每个 .py 文件有明确的单一职责。'
    '最终交付物是通过 build_submission.py 脚本将所有模块按固定顺序拼接成的单个 team_controller.py 文件，'
    '可直接上传至竞赛平台。'
)

# Module table
table = doc.add_table(rows=8, cols=4, style='Light Grid Accent 1')
headers = ['模块文件', '职责', '输入', '输出']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

modules = [
    ('common.py', '数据结构定义', '—', 'PerceptionObs, TrackState, ControlCmd'),
    ('params.py', '参数集中管理', 'profile名', 'CONTROL dict'),
    ('opponent.py', '对手车辆检测', 'BGR图像', 'near_obstacle: bool'),
    ('perception.py', '道路感知', '左右图像+时间戳', 'PerceptionObs'),
    ('estimator.py', '赛道几何估计', 'PerceptionObs+时间戳', 'TrackState'),
    ('policy.py', '控制策略', 'TrackState+时间戳', 'ControlCmd'),
    ('team_controller_local.py', '入口接线', '平台标准输入', '(steering, speed)'),
]
for i, (mod, role, inp, out) in enumerate(modules):
    table.rows[i+1].cells[0].text = mod
    table.rows[i+1].cells[1].text = role
    table.rows[i+1].cells[2].text = inp
    table.rows[i+1].cells[3].text = out

doc.add_paragraph()

doc.add_heading('2.2 关键数据结构', level=2)

add_para(
    '三个核心 dataclass 构成了模块间的通信契约：'
)

add_para(
    'PerceptionObs：包含 center_points（N×2 的 numpy 数组，道路中心采样点）、'
    'left_edge_points / right_edge_points（左右边界点）、road_width_est（道路宽度估计）、'
    'confidence（感知置信度）、line_offset / line_heading / line_confidence（白线状态）、'
    'left_margin_near / right_margin_near（近处边界余量）、near_obstacle（对手车标志）。'
    '所有位置量均为归一化比值。'
)

add_para(
    'TrackState：包含 lateral_error、heading_error、curvature、lookahead_error、confidence 五个核心几何量，'
    '加上 lost（丢线布尔标志）、red_environment（复杂赛道红色环境标志）、'
    '以及从 PerceptionObs 透传的白线和障碍物信息。'
)

add_para(
    'ControlCmd：简单的 steering 和 speed 浮点数对，由策略层产生、入口函数用 clamp_cmd() 做最终裁剪。'
)

doc.add_heading('2.3 构建与测试体系', level=2)

add_para(
    'scripts/build_submission.py 负责将 controller/ 下的模块按固定顺序（common → params → opponent → '
    'perception → estimator → policy → team_controller_local）拼接为自包含的单文件。'
    '构建时自动删除本地 import 语句和调试 I/O（如 open/json/cv2.imwrite），确保提交文件通过官方沙箱校验。'
    '调试构建（--debug-log + --dump-frames）可保留日志和帧输出，仅供本地 Webots 测试使用。'
)

add_para(
    'scripts/validate_submission.py 对本仓库的接口契约、静态规则和 mock 图像输出范围进行本地校验。'
    '官方 SDK 的 validate_controller.py 进行更严格的沙箱合规检查。'
    '全套 pytest 测试（当前 122 个）覆盖了从感知 dropout、车道线跟踪、入弯门控、'
    '脱困策略到输出范围验证的所有关键路径。'
)

# ═══════════════════════════════════════════════════════════════════
# CHAPTER 3: 测试过程报告
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('第3章 测试过程报告', level=1)

doc.add_heading('3.1 单车测试：从无法完赛到稳定巡航', level=2)

add_para(
    '我们的测试始于 basic 椭圆赛道，随后扩展到 complex 复杂赛道。'
    '初期（R001-R010 阶段），核心问题是车辆感知不稳定、入弯太早、以及低速卡死。'
    '通过逐步引入颜色空间多通道分割、草地/护栏检测、白线上下文验证、入弯时机门控和曲率可信度门控，'
    'basic 赛道首先实现了稳定的物理完赛（约 288 秒完成一圈）。'
)

add_para(
    'complex 赛道的问题更加复杂：赛道包含多个复合弯（CP3 区域的发卡弯和道路合并）、'
    '红色地面材质、以及路面上的静止车辆。'
    '关键突破来自 R042-R049 阶段：入弯门控从基于航向偏差（heading）切换到基于近处横向偏差（lateral），'
    '从根本上解决了"直道提前切内线"的问题；配合弯中内侧辅助（inside_assist）和转角迟滞保持（latch），'
    '车辆在 90° 急弯和缓弯中都保持稳定走线。R049 最终实现了 mean 速度 0.85、median 速度 0.90 的巡航性能，'
    'contact 日志显示全程无硬撞栏，仅 7 次轻微擦碰。'
)

doc.add_heading('3.2 多车极端场景测试', level=2)

add_para(
    '多车测试是我们工作的重点，目标是在 6 辆车同时运行的极端条件下，'
    '确保我们的赛车（car_1）能够完成至少一圈。测试在 complex 赛道上以 --batch --fast（无渲染加速）模式运行。'
)

add_para('3.2.1 第一轮测试（基线）', bold=True, size=11)

add_para(
    '6 辆车使用相同的控制器代码同时发车。car_1 在前 280 秒表现良好，速度多次达到 5.80 m/s 的峰值，'
    '但在 t≈283 秒时于 CP3 发卡弯区域（x≈42, y≈146）完全卡死。遥测显示物理速度降至 0.01 m/s，'
    '持续到 300 秒模拟结束。分析发现：CP3 是 6 车拥堵的热点区域，多辆车在此发生碰撞后被挤入栏杆，'
    '而当时的脱困策略（fixed steering 0.74 + high speed 0.90）不但无法救出，高速反而将车辆更紧地推入栏杆。'
)

add_para('3.2.2 第二轮测试（脱困 v1）', bold=True, size=11)

add_para(
    '改进：引入脱困摆动机制（每 10 帧翻转 ±0.20 的舵角）、降低脱困速度（0.90→0.40）、'
    '增加物理卡死检测（speed ≤ 0.06 持续 90 帧直接触发 force_escape）。'
    '结果：car_1 仍于 t≈191 秒在 CP3 卡死。脱困虽触发了摆动，但摆动幅度不够大，'
    '且入弯时未针对多车环境做额外的速度预判。'
)

add_para('3.2.3 第三轮测试（脱困 v2 + 弯道对手减速）— 成功', bold=True, size=11)

add_para(
    '进一步改进：增加弯道+对手车叠加减速（总减速因子 0.40），使车辆在多车拥堵的弯道中以更低速度通过；'
    '摆动幅度增大到 ±0.30；降低脱困置信门槛（0.48→0.25）以适应多车遮挡；缩短物理卡死检测触发时间（90→60 帧）。'
)

add_para(
    '结果：car_1 成功完成了至少 3 圈完整赛道。实时监测显示，车辆在 t=191.0 秒于 CP3 短暂卡死'
    '（speed=0.05），但仅 2 秒后就在 t=192.8 秒恢复到 speed=5.77——摆动脱困成功将车辆从栏杆中"摇出"。'
    '在 t=275.9 秒时，car_1 仍以 5.79 m/s 的速度巡航，全程最高速度达 5.80 m/s。'
)

add_figure(os.path.join(FIG_DIR, 'fig5_run_comparison.png'),
           '图 3.1 三轮测试结果对比')

add_figure(os.path.join(FIG_DIR, 'fig4_multicar_positions.png'),
           '图 3.2 6 车各关键时刻的位置分布')

add_para('3.2.4 多车场景覆盖', bold=True, size=11)

add_para(
    '前车堵路场景：在 6 车拥堵的 CP3 区域，car_1 在检测到近处障碍物后，速度被弯道叠加减速因子降至正常水平的约 40%，'
    '同时主动避让转向（opponent_avoid_steering）根据左右边界余量差朝开阔侧施加偏置。'
    '被撞场景：多车碰撞导致的偏离触发顶栏脱困和低速脱困链。'
    '卡栏杆场景：CP3 的短暂卡死被摆动脱困在 2 秒内救回，证明了 Level 3 脱困机制的有效性。'
)

doc.add_heading('3.3 数据统计分析', level=2)

add_para(
    '图 3.3 展示了第三轮成功测试中 car_1 的速度分布。'
    '在 26,590 个有效遥测帧中（剔除发车后 2 秒内的帧），car_1 的速度分布呈现明显的双峰特征：'
    '主峰集中在 5.5-5.8 m/s（直道和缓弯巡航），次峰在 2.5-4.5 m/s（急弯减速段）。'
    '整体呈右偏分布，说明车辆大部分时间以高速运行。'
)

add_figure(os.path.join(FIG_DIR, 'fig2_speed_dist.png'),
           '图 3.3 car_1 速度分布直方图与时间序列（Run 3, 6-Car Complex）')

add_para(
    '三轮测试的对比数据清晰地展示了策略迭代的效果：第一轮（基线）car_1 在约 283 秒时卡死，'
    '累计卡死时间约 17 秒后仍未恢复。第二轮（脱困 v1）虽然也卡死，但卡死位置更早（191 秒），'
    '说明车辆以更高速度到达 CP3。第三轮（脱困 v2 + 弯道减速）的卡死时间仅为约 2 秒，'
    '随后自动恢复并继续行驶——这证明了摆动脱困和多车弯道减速的组合策略是有效的。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# CHAPTER 4: 小组分工与总结
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('第4章 小组分工与总结', level=1)

doc.add_heading('4.1 小组分工', level=2)

add_para('[请在此处填写小组成员及各自负责的工作内容。]', italic=True)
doc.add_paragraph()

# Team table template
team_table = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
team_headers = ['姓名', '角色', '主要负责模块', '贡献说明']
for i, h in enumerate(team_headers):
    team_table.rows[0].cells[i].text = h
    for p in team_table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
for row_idx in range(1, 4):
    for col_idx in range(4):
        team_table.rows[row_idx].cells[col_idx].text = ''
doc.add_paragraph()

doc.add_heading('4.2 参赛体验', level=2)

add_para('[请在此处撰写参赛体验、心得体会和收获。]', italic=True)
doc.add_paragraph()
add_para('[例如：使用 Webots 仿真平台进行自动驾驶算法开发的感受、]', italic=True)
add_para('[调试视觉算法的挑战与经验、多车协同测试的体会等。]', italic=True)
doc.add_paragraph()

doc.add_heading('4.3 实验总结', level=2)

add_para(
    '本项目的核心目标是为 AI Racer 竞赛平台开发一个鲁棒的视觉自动驾驶控制器，'
    '能够在单车和多车场景下稳定完赛。经过三个阶段的迭代开发，我们取得了以下成果：'
)

add_para('感知层面', bold=True)
add_para(
    '实现了基于多颜色空间联合判据的道路分割，能够区分沥青、草地、红色地面和白色车道线。'
    '车道线检测引入了"路面上下文验证"机制，从根本上排除了护栏支柱等干扰。'
    '跨帧时间锚定（Temporal Anchor）解决了 complex 赛道 CP3 区域道路段合并带来的感知歧义。'
)

add_para('控制策略层面', bold=True)
add_para(
    '引入入弯时机门控（Turn-In Gate），用近处横向偏差而非远处视觉信号作为入弯判断依据，'
    '从根本上解决了直道提前切内线的问题。配合跨帧 latch 保持和出弯迟滞衰减，'
    '使车辆在 90° 急弯和缓弯中均能保持稳定走线。'
    '速度-半径耦合的参考基准自动调整使得高速入弯时门控提前打开，弥补高速冲过的距离。'
)

add_para('多车鲁棒性层面', bold=True)
add_para(
    '建立了四层递进的脱困系统（Pinned → Boundary → Low-Speed → Force Escape），'
    '从依赖精确几何判断到完全独立于几何条件的安全网。'
    '引入摆动脱困策略（每 10 帧 ±0.30 的舵角摆动），模拟人类司机在卡死时的"左右摇轮"操作。'
    '弯道+对手车叠加减速（0.40 总减速因子）和主动避让转向（基于边界余量差的舵角偏置）'
    '使车辆能在 6 车拥堵的复杂赛道上完成多圈巡航。'
)

add_para(
    '实车测试数据表明：在 6 车 complex 赛道的极端测试中，车辆成功完赛 3 圈以上，'
    '仅在 CP3 发卡弯短暂卡死约 2 秒后即自行恢复。全流程完成了 122 个自动化测试用例的覆盖，'
    '以及本地和官方双层校验的通过。'
)

add_para(
    '未来工作方向包括：进一步完善感知一致性（解决相似结构弯道产生不同 curve_risk 估计的问题）、'
    '引入路径规划层以实现主动超车策略、以及将侧方/后方来车检测纳入对手感知模块。'
)

# ── Save ──────────────────────────────────────────────────────────
output_path = os.path.join(REPO_ROOT, 'experiments', 'AI_Racer_Experiment_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
